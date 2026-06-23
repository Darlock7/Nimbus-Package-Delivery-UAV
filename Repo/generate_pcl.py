"""
NIMBUS PCL Rev 4.0 — updated for final configuration.
Changes from Rev 3:
- Phugoid confirmed stable (zeta=0.058, Level 1) per latest run
- Prop name: APC 10x4.7 (no type suffix)
- Static thrust updated to 12.208 N (Rm=0.148 ohm, latest surrogate run)
- Ground roll updated to 34.1 m (104.3 m margin)
- Prop clearance corrected to 73 mm
- Rudders REMOVED — steerable nose gear replaces ground yaw
- New Section 2: Controls & Channel Map
- Pre-flight checklist updated (no rudder checks, added nose gear steering check)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

NAVY = RGBColor(0x1A, 0x37, 0x6E)
RED  = RGBColor(0xC0, 0x00, 0x00)
GRN  = RGBColor(0x1D, 0x6A, 0x27)
WHT  = RGBColor(0xFF, 0xFF, 0xFF)
GRY  = RGBColor(0x55, 0x55, 0x55)

# ── helpers ─────────────────────────────────────────────────────────────────
def cell_bg(cell, hex6):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), hex6)
    pr.append(s)

def hline(color='1A376E', sz='4'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    pr = p._p.get_or_add_pPr()
    bd = OxmlElement('w:pBdr')
    b  = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
    b.set(qn('w:space'), '1');    b.set(qn('w:color'), color)
    bd.append(b); pr.append(bd)

def sec(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(txt.upper())
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY

def sub(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = NAVY

def body(txt, size=10, bold=False, color=None, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(txt)
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color

def warn(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run('⚠  ' + txt)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RED

def note(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run('ℹ  ' + txt)
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = NAVY

def check(txt, size=10):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run('☐  ' + txt); r.font.size = Pt(size)

def num(bold_part, rest=''):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(bold_part); r1.bold = True; r1.font.size = Pt(10)
    if rest:
        r2 = p.add_run('  ' + rest); r2.font.size = Pt(10)

def tbl(rows, col_widths=None):
    n = len(rows[0])
    t = doc.add_table(rows=len(rows), cols=n)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            if ri == 0:
                cell_bg(cell, '1A376E')
                r = cell.paragraphs[0].add_run(str(val))
                r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHT
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell_bg(cell, 'F5F5F5' if ri % 2 == 1 else 'FFFFFF')
                r = cell.paragraphs[0].add_run(str(val))
                r.font.size = Pt(9.5)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for ri in range(len(rows)):
            for ci, w in enumerate(col_widths):
                t.rows[ri].cells[ci].width = Cm(w)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NIMBUS  ·  FLY DAY CHECKLIST')
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('MAE 155B · Group 2 · Flying Wing RC Aircraft · Rev 4.0 · 2026-06-11')
r2.font.size = Pt(9); r2.font.color.rgb = GRY

hline(sz='8')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — AIRCRAFT LIMITS
# ══════════════════════════════════════════════════════════════════════════════
sec('1. Aircraft Limits')

tbl([
    ['Parameter',                  'Value',        'Imperial'],
    ['Max Gross Weight (loaded)',   '2.60 kg',      '5.73 lb'],
    ['Payload (volume box + pkg)',  '1.10 kg total','2.43 lb'],
    ['Stall Speed  V_S',           '12.9 m/s',     '28.8 mph'],
    ['Takeoff Speed  V_TO',        '14.4 m/s',     '32.2 mph'],
    ['Approach Speed  (1.2 × V_S)','15.5 m/s',     '34.7 mph'],
    ['Design Cruise  V_C',         '20 m/s',       '44.7 mph'],
    ['Never-Exceed  V_NE',         '25.0 m/s',     '56.0 mph'],
    ['Maneuver Speed  V_A',        '24.1 m/s',     '53.9 mph'],
    ['Best Glide Speed',           '~18 m/s',      '~40 mph'],
    ['Max Bank Angle',             '51.2 deg',     '—'],
    ['Min Turn Radius (20 m/s)',   '36.5 m',       '119.7 ft'],
    ['Max Load Factor  (+)',       '+3.8 g',       '—'],
    ['Max Load Factor  (−)',       '−1.5 g',       '—'],
], col_widths=[7.0, 4.5, 4.0])

warn('Do NOT exceed 25.0 m/s (V_NE). Stay at 20–24 m/s in cruise.')

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CONTROLS & CHANNEL MAP
# ══════════════════════════════════════════════════════════════════════════════
sec('2. Controls & Transmitter Channel Map')

note('NIMBUS has NO rudder surfaces. The vertical fins are fixed. '
     'Yaw control on the ground comes from the steerable nose gear only. '
     'In the air, all maneuvering is done with elevons (bank to turn).')

sub('Control Surfaces & Actuators')
tbl([
    ['What',                   'How it works',                                   'When active'],
    ['Elevons (left & right)', 'Symmetric deflection = pitch.  Differential = roll.  Elevon mixing ON in TX.', 'IN FLIGHT'],
    ['Steerable nose gear',    'Servo steers the nose wheel left/right.  Gives directional control on the runway.', 'GROUND ONLY'],
    ['Cargo bay door',         'Servo opens the bottom door to release the 300 g volume box.', 'RELEASE ONLY'],
    ['Motor kill switch',      'Cuts motor signal to ESC independently of throttle stick.', 'SAFETY / EMERGENCY'],
], col_widths=[4.0, 8.0, 4.0])

sub('Mode 2 Transmitter Channel Map')
note('Mode 2: left stick = throttle (up/down) + nose gear steering (left/right). '
     'Right stick = pitch (up/down) + roll (left/right). '
     'Elevon mixing must be enabled in the transmitter model file.')

tbl([
    ['Channel', 'TX Control',              'Aircraft Function',                   'Notes'],
    ['Ch 1',    'Right stick — LEFT/RIGHT','Elevon differential  →  ROLL',        'Aileron input into mixer'],
    ['Ch 2',    'Right stick — UP/DOWN',   'Elevon symmetric  →  PITCH',          'Elevator input into mixer'],
    ['Ch 3',    'Left stick  — UP/DOWN',   'Throttle  →  Motor / ESC',            'Idle at bottom before arming'],
    ['Ch 4',    'Left stick  — LEFT/RIGHT','Steerable nose gear  →  YAW (ground)','No effect once airborne'],
    ['Ch 5',    'Switch (e.g. Aux 1)',     'Cargo bay door  →  Payload release',  'Two-position: closed / open'],
    ['Ch 6',    'Switch (e.g. Aux 2)',     'Motor kill switch',                   'Safety cut — disables motor'],
], col_widths=[1.8, 4.2, 5.0, 5.5])

sub('Elevon Deflection Logic  (confirm with team member during preflight)')
tbl([
    ['Pilot Input',         'Left Elevon',   'Right Elevon',  'Aircraft Response'],
    ['Right stick UP',      'TE UP',         'TE UP',         'Pitch nose UP'],
    ['Right stick DOWN',    'TE DOWN',       'TE DOWN',       'Pitch nose DOWN'],
    ['Right stick RIGHT',   'TE DOWN',       'TE UP',         'Roll RIGHT (bank right)'],
    ['Right stick LEFT',    'TE UP',         'TE DOWN',       'Roll LEFT (bank left)'],
], col_widths=[3.8, 3.0, 3.0, 5.7])

warn('No rudder = no yaw authority in flight. All turns are banked turns via roll. '
     'Do not expect coordinated yaw — the aircraft self-coordinates as a flying wing.')

warn('Nose gear steering (Ch 4) is ONLY effective while the nose wheel is on the ground. '
     'Do not use it in flight.')

warn('Cargo bay door (Ch 5): confirm switch is in CLOSED position before every flight. '
     'Accidentally opening on takeoff will shed the payload.')

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PROPULSION SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
sec('3. Propulsion System')

tbl([
    ['Item',              'Spec',                      'Limit / Note'],
    ['Motor',             'SunnySky 2212 — 1100 KV',   'Max continuous 34 A'],
    ['Battery',           '3S LiPo — 11.1 V nominal',  'Min pre-flight 11.7 V  (3.9 V/cell)'],
    ['Propeller',         'APC 10×4.7',                'No nicks or cracks — inspect before every flight'],
    ['Static Thrust',     '12.2 N  (1.24 kgf)',        'At 11.1 V, Rm = 0.148 Ω — surrogate model'],
    ['Static Current',    '27.4 A',                    '34 A limit — OK (80% of limit)'],
    ['Batt min in-flight','10.5 V  (3.5 V/cell)',      'LAND IMMEDIATELY if reached'],
], col_widths=[4.0, 6.0, 6.5])

note('Static thrust 12.2 N vs aircraft weight 25.5 N gives T/W = 0.48 — adequate for takeoff and climb.')

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WEIGHT & BALANCE
# ══════════════════════════════════════════════════════════════════════════════
sec('4. Weight & Balance')
warn('CRITICAL — Verify CG before EVERY flight.  SM = 5.5%  (low margin). '
     'A misplaced battery or payload WILL make the aircraft uncontrollable.')

tbl([
    ['State',                          'Mass',     'CG from Nose',        '% MAC'],
    ['State 1 — Loaded (box + pkg)',   '2.60 kg',  '348 mm  (13.70 in)',  '22.5%'],
    ['State 2 — Box dropped (pkg only)','2.30 kg', '348 mm  (13.70 in)',  '22.5%  (no CG shift)'],
    ['State 3 — Ferry / empty',        '1.50 kg',  '348 mm  (13.69 in)',  '22.4%'],
    ['AFT LIMIT  (Neutral Point)',      '—',        '353 mm  (13.90 in)',  '25.0%  ← NEVER EXCEED'],
], col_widths=[5.5, 2.5, 5.0, 3.5])

note('CG stays at 348 mm after the volume box drops because both payloads sit at the same x-station. '
     'No trim change needed after payload release.')

sub('CG Check Procedure')
num('Fully assemble', '— battery in, payload secured, prop on.')
num('Lift aircraft', 'with two fingers at 348 mm (13.70 in) from nose.')
num('Must balance', 'level ± 2 deg.')
num('Nose-heavy', '→ move battery AFT.   Tail-heavy → move battery FORWARD.')
num('Re-check', 'after any component swap or battery change.')

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — PILOT NOTES
# ══════════════════════════════════════════════════════════════════════════════
sec('5. Pilot Notes — Flight Characteristics')

tbl([
    ['Mode',          'Period / τ',   'Damping ζ',              'Pilot Action'],
    ['Short Period',  '0.90 s',       '0.61 — well damped',     'Normal pitch inputs — crisp response'],
    ['Phugoid',       '~16 s',        '0.058 — lightly damped', 'One small input per cycle peak — do NOT chase the oscillation'],
    ['Dutch Roll',    '1.2 s cycle',  '0.082 — Level 2',        'Release inputs and let it damp — no rudder available'],
    ['Roll Subsid.',  'τ = 0.042 s',  'Level 1 — very crisp',   'Brief stick inputs only — overshoot likely with large inputs'],
    ['Spiral',        't×2 = 20.9 s', 'Level 1 — slow',         'Check and level the bank every 20 s'],
], col_widths=[3.0, 2.8, 4.0, 6.2])

warn('Dutch Roll cannot be damped with rudder — there is no rudder. Release all inputs and wait.')
body('Trim: ~5 deg elevon TE down (pre-trimmed).  Pitches UP → more TE down.  Pitches DOWN → less TE down.',
     size=9.5, indent=0.3)

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PRE-FLIGHT CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════
sec('6. Pre-Flight Checklist')
warn('Complete ALL sections in order. Do not proceed to runway until Section G is signed off.')

sub('A — Vehicle Inspection')
for item in [
    'No cracks, delamination, or damage on wing surfaces',
    'Spar tube (10 mm) fully seated through both wing panels — no play',
    'Wing joint secure — no rocking',
    'Both elevons: move freely, full travel, hinges intact, no binding',
    'Vertical fins: secure, no wobble (fins are FIXED — no rudder servos)',
    'Prop: APC 10×4.7 — no nicks, cracks — firmly seated, spinner tight',
    'Motor: spin by hand → smooth, no grinding, no wobble',
    'All motor mount screws tight',
    'Landing gear: all 3 legs secure and straight',
    'Nose gear: steers cleanly left and right by hand, returns to center',
    'Payload (volume box + package) secured — no movement inside cargo bay',
    'Cargo bay door: fully closed and latched in pre-release position',
]:
    check(item)

sub('B — Battery & Electrical')
warn('Battery must be ≥ 11.7 V before flight. DO NOT fly if below this.')
for item in [
    'Battery voltage ≥ 11.7 V  (≥ 3.9 V/cell)                  Measured: _______ V',
    'Battery polarity confirmed before plugging in',
    'Battery strap tight — cannot shift in flight',
    'ESC wiring clear of prop arc and control linkages',
    'Receiver seated, antenna oriented correctly',
    'All 6 servo/ESC connectors fully seated in receiver',
]:
    check(item)

sub('C — Transmitter Setup')
warn('POWER ON TRANSMITTER BEFORE connecting aircraft battery.')
for item in [
    'Transmitter ON first — before battery connect',
    'Model file: NIMBUS selected',
    'Elevon mixing: ON  (both elevons driven from elevator + aileron inputs)',
    'Elevon throw: ± 20 deg — verify at full stick deflection',
    'Ch 4 (nose gear): responds to left stick left/right',
    'Ch 5 (cargo door): switch in CLOSED / safe position',
    'Ch 6 (motor kill): switch in ARMED / normal position',
    'Throttle stick: full idle (bottom) before battery connect',
    'All trims centered or saved NIMBUS trim file loaded',
    'Failsafe set: throttle → idle on signal loss',
]:
    check(item)

sub('D — Battery Connect & ESC Check')
for item in [
    'Call "CONNECTING BATTERY" — all personnel clear of prop arc',
    'Connect battery — listen for ESC initialization beep sequence',
    'Wait 5 seconds — do NOT touch throttle',
    'Advance throttle to 10% — motor spins up smoothly (no grinding)',
    'Return to idle — no error tones',
]:
    check(item)

sub('E — Control Surface & Actuator Function Check')
note('A second team member confirms deflections and actuator movement at the aircraft.')
for item in [
    'Right stick UP   →  both elevons: trailing edge UP  ✓',
    'Right stick DOWN →  both elevons: trailing edge DOWN  ✓',
    'Right stick RIGHT →  right elevon DOWN, left elevon UP  ✓',
    'Right stick LEFT  →  left elevon DOWN, right elevon UP  ✓',
    'Left stick LEFT/RIGHT →  nose gear steers left/right smoothly  ✓',
    'No binding at full deflection on any elevon  ✓',
    'Ch 5 switch → cargo door opens smoothly, closes fully  ✓  (confirm CLOSED after test)',
    'Ch 6 switch → motor kill cuts motor instantly, ARMED restores  ✓',
]:
    check(item)

sub('F — Range Check')
for item in [
    'Carry aircraft 30 m from transmitter (battery connected)',
    'Cycle all controls through full deflection — no glitching or lag',
    'Throttle test at 30 m — motor responds cleanly',
    'Return aircraft to launch point',
]:
    check(item)

sub('G — Weight & Balance — Final Sign-Off')
warn('DO NOT move to runway until this section is complete.')
for item in [
    'CG balances at 348 mm from nose  ± 5 mm',
    'Total payload confirmed: volume box + package              Weighed: _______ g',
    'Cargo door switch confirmed in CLOSED position',
    'Battery locked in position',
    'CG is NOT aft of 353 mm from nose',
    'All team members confirm: GO for flight',
]:
    check(item)

hline()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — TAKEOFF
# ══════════════════════════════════════════════════════════════════════════════
sec('7. Takeoff Procedure')
warn('Point aircraft DIRECTLY into wind. Use nose gear steering (left stick) to hold centerline during roll.')

tbl([
    ['Takeoff Parameter',       'Value'],
    ['Takeoff speed  V_TO',     '14.4 m/s  (32.2 mph)'],
    ['Ground roll distance',    '34.1 m  (112 ft)'],
    ['Available runway',        '138.4 m  (454 ft)'],
    ['Runway margin',           '104.3 m  (342 ft) remaining after liftoff'],
    ['Static thrust',           '12.2 N  — T/W = 0.48  (adequate for 6° climb)'],
    ['Prop clearance',          '73 mm  (2.9 in) above ground'],
    ['Ground incidence',        '3.2 deg nose-up — gear does most of rotation'],
], col_widths=[6.5, 9.0])

note('Ground roll computed with Nicolai slide method: V_eval = 10.1 m/s, '
     'T = 9.17 N at 0.7×VTO, friction μ = 0.03, a_mean = 3.04 m/s².')

sub('Takeoff Steps')
num('Point into wind.', '')
num('Advance throttle to FULL — smoothly.', 'Use left stick (Ch 4) to track centerline with nose gear.')
num('Aircraft lifts off at ~14.4 m/s.', 'Ground roll ≈ 34 m.  Nose gear steering has NO effect once airborne.')
num('Light back-stick ONLY at liftoff.', 'Gear gives 3.2° nose-up — only 2.5° more needed.  A hard pull will stall.')
num('Maintain +6 deg climb at 11 m/s.', '')
num('Climb out smoothly.', 'Do not aggressively chase altitude.')

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — MISSION & PAYLOAD RELEASE
# ══════════════════════════════════════════════════════════════════════════════
sec('8. Mission & Payload Release')

note('State 1 → State 2: flip Ch 5 switch to open cargo door.  Volume box (300 g) drops.  '
     'CG stays at 348 mm — no trim change needed after release.')

for item in [
    'Climb to mission altitude and establish cruise (20 m/s)',
    'Complete required laps / fly to delivery point',
    'Over delivery point: flip Ch 5 switch → open cargo door → volume box releases',
    'Confirm door opened (visual from spotter)',
    'Return Ch 5 to closed to reset for landing inspection',
    'Continue mission or proceed to landing',
]:
    check(item)

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — CRUISE
# ══════════════════════════════════════════════════════════════════════════════
sec('9. Cruise')

tbl([
    ['Parameter',                'Value'],
    ['Target cruise speed',      '20–24 m/s  (45–54 mph)'],
    ['HARD ceiling  V_NE',       '25.0 m/s  (56 mph)  — do NOT exceed'],
    ['Elevon trim',              '~5 deg TE down (pre-trimmed)'],
    ['Max bank',                 '51.2 deg  — all turns via bank, no rudder'],
    ['Min turn radius',          '36.5 m  — do not tighten at mission speed'],
    ['Battery — land if below',  '10.5 V  (3.5 V/cell)'],
], col_widths=[6.5, 9.0])

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — LANDING
# ══════════════════════════════════════════════════════════════════════════════
sec('10. Landing Procedure')

num('Approach at 15.5 m/s (35 mph)', '— 1.2 × V_stall.')
num('Throttle to idle', '— maintain approach angle with elevator (right stick).')
num('Maintain runway heading with bank', '— small bank corrections only, no rudder available.')
num('At 1–2 m AGL: gentle back-stick flare', '— bleed to ~14 m/s.')
num('Main gear touches first', 'in 3.2° nose-up attitude.')
num('After main gear contact: elevons NEUTRAL', '— let nose gear settle gently.')
num('Once on the ground:', 'use nose gear steering (left stick) to track centerline during rollout.')
warn('Do NOT hold nose up after touchdown — prop has 73 mm clearance. Let the nose down promptly.')

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — EMERGENCY PROCEDURES
# ══════════════════════════════════════════════════════════════════════════════
sec('11. Emergency Procedures')

sub('Motor Failure')
num('Hold ~18 m/s', '(best glide, L/D ≈ 14).  Do NOT pull back hard.')
num('Turn toward open landing area', '— use bank (roll), no rudder available.')
num('Normal approach at 15.5 m/s', ', flare at 1–2 m AGL.')

sub('Runaway Motor — Use Motor Kill Switch')
num('Flip Ch 6 kill switch', '— motor cuts immediately.')
num('Glide to landing', '— treat as motor failure from this point.')
warn('Motor kill switch (Ch 6) is the primary safety tool — practice its location before flight day.')

sub('Phugoid — Slow Pitch Oscillation (~16 s cycle)')
num('Identify:', 'nose slowly rising and falling every ~16 seconds.')
num('DO NOT CHASE', '— wait for pitch to slow at the top or bottom of the cycle.')
num('One small smooth elevator input per cycle peak.', '2–3 corrections will damp it.')

sub('Dutch Roll — Yaw/Roll Coupling')
num('Release ALL inputs', '— allow to self-damp.  No rudder available to help.')
num('If it continues:', 'small roll input opposite to the rolling side only.')
num('Avoid large aileron inputs', '— they will excite dutch roll further.')

sub('Spiral Dive — Bank Slowly Steepening')
num('Aileron toward raised wing', '— level the wings.')
num('Resume cruise', '— spiral takes ~21 s to become critical.')

hline()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — ABORT CRITERIA
# ══════════════════════════════════════════════════════════════════════════════
sec('12. Abort Criteria — DO NOT FLY IF:')
warn('ANY single condition below = ground the aircraft.  Fix it before retrying.')

for item in [
    'CG not balancing at 348 mm ± 5 mm from nose',
    'Battery below 11.7 V pre-flight',
    'Any elevon reversed, binding, or not reaching full throw',
    'Nose gear not steering — no directional control on ground',
    'Prop damage: any nick, crack, or wobble',
    'Payload not secured or cargo door not latching closed',
    'Wind above team-agreed limit (suggest ≤ 7 m/s)',
    'Any ESC error tone or motor irregularity',
    'Range check failed: glitching at 30 m',
    'CG aft of 353 mm — at or past neutral point',
    'Any structural damage: wing cracks, loose spar, delamination',
    'Motor kill switch not functioning',
    'Ch 5 cargo door switch cannot be confirmed in CLOSED position',
]:
    check(item)

hline()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — QUICK REFERENCE CARD
# ══════════════════════════════════════════════════════════════════════════════
sec('13. Quick Reference — Cut Out & Keep')

lines = [
    'NIMBUS  ·  MAE 155B Group 2  ·  Rev 4.0  (2026-06-11)',
    '─' * 56,
    'WEIGHTS',
    '  MTOW (loaded)    2.60 kg     5.73 lb',
    '  Payload total    1.10 kg     (300g box + 800g pkg)',
    '  Empty            1.50 kg     3.31 lb',
    '',
    'SPEEDS',
    '  V_stall         12.9 m/s    28.8 mph',
    '  V_TO            14.4 m/s    32.2 mph',
    '  V_cruise        20 m/s      44.7 mph',
    '  V_NE            25.0 m/s    56 mph   ← HARD LIMIT',
    '  V_approach      15.5 m/s    34.7 mph  (1.2 Vs)',
    '  V_best_glide    ~18 m/s     ~40 mph',
    '',
    'TAKEOFF',
    '  Ground roll      34.1 m     112 ft',
    '  Runway avail    138.4 m     454 ft',
    '  Runway margin   104.3 m     342 ft',
    '',
    'CG & STABILITY',
    '  CG (loaded)      348 mm     13.70 in from nose',
    '  AFT LIMIT        353 mm     13.90 in  ← DO NOT EXCEED',
    '  Static margin    5.5 %      (target 5-10 %)',
    '',
    'CHANNEL MAP',
    '  Ch 1  Right stick L/R  →  Roll (elevon diff)',
    '  Ch 2  Right stick U/D  →  Pitch (elevon sym)',
    '  Ch 3  Left stick  U/D  →  Throttle',
    '  Ch 4  Left stick  L/R  →  Nose gear (GROUND ONLY)',
    '  Ch 5  Switch Aux 1     →  Cargo door (payload drop)',
    '  Ch 6  Switch Aux 2     →  Motor kill switch',
    '',
    'PROPULSION',
    '  Motor   SunnySky 2212 — 1100 KV',
    '  Battery 3S LiPo — 11.1 V nominal',
    '  Prop    APC 10x4.7',
    '  Static thrust  12.2 N  (27.4 A static)',
    '  Batt min pre-flight  11.7 V  (3.9 V/cell)',
    '  Batt min in-flight   10.5 V  (3.5 V/cell)  -> LAND NOW',
    '',
    'CONTROL THROWS',
    '  Elevon    +/- 20 deg',
    '  Trim      +5 deg TE down',
    '  Max bank  51.2 deg',
    '  Min turn  36.5 m  (at 20 m/s)',
    '─' * 56,
]

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after  = Pt(2)
for line in lines:
    r = p.add_run(line + '\n')
    r.font.name = 'Courier New'; r.font.size = Pt(9)

hline()
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_foot.add_run(
    'Rev 4.0  ·  2026-06-11  ·  MAE 155B Group 2  ·  '
    'Verify against latest outputs/manufacturing_dimensions.txt before flight day'
)
rf.font.size = Pt(7.5); rf.italic = True; rf.font.color.rgb = GRY

out = ('/Users/JuanManuelSanchez/Desktop/MAE155B-Aircraft-Design/'
       'class_material/Reports&Presentations/NIMBUS_PCL_Rev4.docx')
doc.save(out)
print(f'Saved -> {out}')
